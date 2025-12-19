"""
DOCX Resume Generator
Converts Markdown resume content to formatted Word documents.
"""
import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import logging

logger = logging.getLogger(__name__)

class ResumeParser:
    def __init__(self, markdown_text):
        self.markdown = markdown_text
        self.sections = {}
        self.parse()

    def parse(self):
        lines = self.markdown.split('\n')
        current_section = None
        content_lines = []
        
        def save_section():
            if current_section and content_lines:
                self.sections[current_section] = '\n'.join(content_lines).strip()

        # Extract name (first line with #)
        name_match = re.search(r'^#\s+(.+)', self.markdown, re.MULTILINE)
        self.sections['name'] = name_match.group(1).strip() if name_match else 'Your Name'

        # Extract title (first bold text **)
        title_match = re.search(r'\*\*(.+?)\*\*', self.markdown)
        self.sections['title'] = title_match.group(1).strip() if title_match else ''

        # Extract contact info
        self.sections['contact'] = self.extract_contact()

        for line in lines:
            line = line.strip()
            if line.startswith('## '):
                save_section()
                raw_name = line[3:].strip().lower()
                current_section = re.sub(r'\s+', '_', raw_name)
                content_lines = []
            elif current_section:
                content_lines.append(line)
        
        save_section()

    def extract_contact(self):
        contact = {}
        email_match = re.search(r'([a-zA-Z0-9._-]+@[a-zA-Z0-9._-]+\.[a-zA-Z0-9_-]+)', self.markdown)
        if email_match: contact['email'] = email_match.group(1)
            
        phone_match = re.search(r'(\+?\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9})', self.markdown)
        if phone_match: contact['phone'] = phone_match.group(1)
            
        linkedin_match = re.search(r'(linkedin\.com\/in\/[^\s)]+)', self.markdown)
        if linkedin_match: contact['linkedin'] = linkedin_match.group(1)
            
        location_match = re.search(r'([A-Z][a-z]+,\s*[A-Z]{2}|[A-Z][a-z]+,\s*[A-Z][a-z]+)', self.markdown)
        if location_match: contact['location'] = location_match.group(1)
            
        return contact
        
    def get_section(self, keys):
        if isinstance(keys, str): keys = [keys]
        for key in keys:
            if key in self.sections: return self.sections[key]
        return ""

def markdown_to_docx(markdown_text, template_id='modern'):
    document = Document()
    parser = ResumeParser(markdown_text)
    
    # Configure Styles
    styles = {
        'template-1': {'header_color': RGBColor(31, 42, 51), 'font': 'Arial', 'sidebar_bg': True}, # Dark Sidebar
        'template-2': {'header_color': RGBColor(0, 0, 0), 'font': 'Calibri', 'accent': RGBColor(220, 20, 60)}, # Split
        'template-3': {'header_color': RGBColor(51, 51, 51), 'font': 'Times New Roman'}, # Classic
        'template-4': {'header_color': RGBColor(0, 77, 64), 'font': 'Calibri', 'sidebar_bg': True}, # Teal
        'template-5': {'header_color': RGBColor(11, 78, 162), 'font': 'Arial'} # Blue
    }
    current_style = styles.get(template_id, styles['template-1'])
    
    # Helper to add markdown content
    def add_markdown_content(container, text, bold_color=None):
        if not text: return
        for line in text.split('\n'):
            line = line.strip()
            if not line: continue
            
            p = container.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            
            if line.startswith('### '):
                run = p.add_run(line[4:])
                run.bold = True
                run.font.size = Pt(11)
            elif line.startswith('- ') or line.startswith('* '):
                p.style = 'List Bullet'
                content = line[2:]
                # Basic bold parsing
                parts = re.split(r'(\*\*.*?\*\*)', content)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        r = p.add_run(part[2:-2])
                        r.bold = True
                        if bold_color: r.font.color.rgb = bold_color
                    else:
                        p.add_run(part)
            else:
                 # Basic bold parsing
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        r = p.add_run(part[2:-2])
                        r.bold = True
                        if bold_color: r.font.color.rgb = bold_color
                    else:
                        p.add_run(part)

    def add_section_header(container, text, color=None):
        p = container.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        if color: run.font.color.rgb = color
        
        # Bottom border mimic (underscore? or just bold)
        # Word borders are hard XML work, simpler to just bold/color for now
        
    # --- Generation Logic ---
    
    # 1. Header (Name/Title) common for most, but T2/T5 have bands logic which is hard in Word.
    # We will use simple blocks.
    
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = header.add_run(parser.sections['name'] + "\n")
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = current_style['header_color']
    
    title_run = header.add_run(parser.sections['title'])
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(100, 100, 100)

    # 2. Layouts
    
    sections = parser.sections
    summary = parser.get_section(['summary', 'professional_summary'])
    skills = parser.get_section(['skills', 'technical_skills'])
    experience = parser.get_section(['experience', 'work_experience'])
    education = parser.get_section(['education'])
    contact = parser.sections['contact']

    contact_text = '\n'.join([v for v in contact.values() if v])

    if template_id in ['template-1', 'template-2', 'template-4', 'template-5']:
        # Multi-column: Use Table
        table = document.add_table(rows=1, cols=2)
        table.autofit = False
        
        # Set Widths (approximate for A4)
        # Col 1: 30%, Col 2: 70% (except T2/T5 which might be different)
        width_1 = Inches(2.3)
        width_2 = Inches(5.0) # Approx remaining
        
        table.columns[0].width = width_1
        table.columns[1].width = width_2
        
        # Sidebar Cell
        cell_sidebar = table.cell(0, 0)
        cell_main = table.cell(0, 1)
        
        # If styled like T2 (left main, right sidebar) or T5
        if template_id in ['template-2', 'template-5']:
             # Swap width logic?
             table.columns[0].width = width_2
             table.columns[1].width = width_1
             cell_main = table.cell(0, 0)
             cell_sidebar = table.cell(0, 1)
        
        # Fill Sidebar
        add_section_header(cell_sidebar, "Contact", current_style['header_color'])
        cell_sidebar.add_paragraph(contact_text)
        
        if skills:
            add_section_header(cell_sidebar, "Skills", current_style['header_color'])
            add_markdown_content(cell_sidebar, skills)
            
        if education:
            add_section_header(cell_sidebar, "Education", current_style['header_color'])
            add_markdown_content(cell_sidebar, education)
            
        # Fill Main
        if summary:
            add_section_header(cell_main, "Profile", current_style['header_color'])
            add_markdown_content(cell_main, summary)
            
        if experience:
            add_section_header(cell_main, "Experience", current_style['header_color'])
            add_markdown_content(cell_main, experience)
            
    else:
        # Single Column (Template 3)
        # Contact
        p = document.add_paragraph(contact_text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if summary:
            add_section_header(document, "Profile", current_style['header_color'])
            add_markdown_content(document, summary)
            
        if skills:
            add_section_header(document, "Skills", current_style['header_color'])
            add_markdown_content(document, skills)
            
        if experience:
            add_section_header(document, "Experience", current_style['header_color'])
            add_markdown_content(document, experience)
            
        if education:
            add_section_header(document, "Education", current_style['header_color'])
            add_markdown_content(document, education)
            
    # Save
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
